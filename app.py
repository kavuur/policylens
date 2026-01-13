# app.py  (first part)
import os
import json
import mimetypes
import hashlib
from datetime import datetime, timedelta
from functools import wraps
import re
import unicodedata
import logging
import sys
import uuid

import requests
import pandas as pd
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from dotenv import load_dotenv
from flask import (
    Flask, render_template, request, redirect, url_for, flash,
    jsonify, send_from_directory, send_file, abort, g, has_request_context
)
from flask_login import (
    LoginManager, current_user, login_user, login_required, logout_user
)
from flask_migrate import Migrate
from flask_mail import Mail, Message
from sqlalchemy.orm import joinedload
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
from wtforms import ValidationError
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_wtf.csrf import CSRFProtect

# --- Models / Services ---
from models.models import (
    db, User, PolicyDocument, Codebook, ResearchNote, Project,
    Media, Descriptor, Code, SubCode, SubSubCode, Excerpt
)
from forms import (
    ProjectForm, ProfileUpdateForm, PasswordUpdateForm,
    CodebookForm, CodebookEditForm, CodeForm, MediaUploadForm
)
from models.document import DocumentProcessor
from services.analysis import AnalysisService
from services.policy_analyzer import PolicyAnalyzer

# -----------------------------------------------------------------------------
# Env & App setup
# -----------------------------------------------------------------------------
load_dotenv()  # ensure env is loaded even under gunicorn

app = Flask(__name__)

# Stable secret (do NOT use os.urandom in prod)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-unsafe-secret')
app.config['WTF_CSRF_ENABLED'] = True
app.config['WTF_CSRF_SECRET_KEY'] = 'another-secret-key'  # Change this to a different secure key

# Load config from config.py
from config import DevelopmentConfig, ProductionConfig
app.config.from_object(DevelopmentConfig if app.debug else ProductionConfig)

mail = Mail(app)

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Trust proxy headers so request.is_secure mirrors external HTTPS
# (works behind Nginx/ALB/Traefik and in local dev)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1)

# Base cookie settings; Secure flags are set per-request below
app.config.update(
    SESSION_COOKIE_SAMESITE='Lax',
    REMEMBER_COOKIE_SAMESITE='Lax',
    REMEMBER_COOKIE_DURATION=timedelta(days=30),
    # If you serve the app under a subpath behind a proxy, leave PATH at '/'
    SESSION_COOKIE_PATH='/',
    # If you need a fixed domain (e.g., app.example.org), set via env:
    SESSION_COOKIE_DOMAIN=os.getenv('SESSION_COOKIE_DOMAIN', None),
)

# Adapt cookie security automatically for HTTP (dev) vs HTTPS (prod)
@app.before_request
def _adapt_cookie_security():
    is_https = request.is_secure  # accurate after ProxyFix
    app.config['SESSION_COOKIE_SECURE'] = is_https
    app.config['REMEMBER_COOKIE_SECURE'] = is_https

# -----------------------------------------------------------------------------
# Database
# -----------------------------------------------------------------------------
db_url = os.getenv('DATABASE_URL')
if not db_url:
    raise RuntimeError('DATABASE_URL must be set to point to your Postgres database')
app.config['SQLALCHEMY_DATABASE_URI'] = db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)
migrate = Migrate(app, db)

# -----------------------------------------------------------------------------
# Uploads  ✅ resilient & writable
# -----------------------------------------------------------------------------
def _ensure_writable_dir(candidates):
    """Return the first writable dir from candidates; create if needed."""
    for path in candidates:
        if not path:
            continue
        try:
            os.makedirs(path, exist_ok=True)
            test_file = os.path.join(path, ".write_test")
            with open(test_file, "wb") as fh:
                fh.write(b"ok")
            os.remove(test_file)
            return path
        except Exception as e:
            app.logger.warning("Upload dir not writable: %s (%s)", path, e)
    raise RuntimeError("No writable upload directory found")

# Prefer env-provided root (e.g., mounted volume) → instance/uploads → /tmp fallback
#UPLOAD_ROOT = os.getenv('UPLOAD_ROOT', os.path.join(app.root_path, 'uploads'))
#app.config['UPLOAD_FOLDER'] = UPLOAD_ROOT
#os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)  # best-effort

import os
from pathlib import Path

# --- Upload directory resolution (robust, non-fatal) ---
# 1) Prefer explicit env UPLOAD_ROOT (e.g., /data/uploads from docker-compose)
# 2) Fall back to instance/uploads (Flask's writable area)
# 3) Last resort: /tmp/policylens_uploads
UPLOAD_ROOT_ENV = os.getenv("UPLOAD_ROOT")  # <-- define it first

# Ensure Flask's instance path exists (Flask may not have created it yet)
INSTANCE_UPLOADS = str(Path(getattr(app, "instance_path", "./instance")).joinpath("uploads"))

candidates = [
    UPLOAD_ROOT_ENV,          # e.g., /data/uploads (may be None)
    "/data/uploads",          # typical Docker volume mount
    INSTANCE_UPLOADS,         # ./instance/uploads
    "/tmp/policylens_uploads" # last resort
]

chosen = None
for d in filter(None, candidates):
    try:
        Path(d).mkdir(parents=True, exist_ok=True)
        chosen = d
        break
    except Exception as e:
        app.logger.warning("Upload dir candidate not usable: %s (%r)", d, e)

if not chosen:
    # Absolute worst case—shouldn't happen because /tmp should work.
    chosen = "/tmp/policylens_uploads"
    Path(chosen).mkdir(parents=True, exist_ok=True)

app.config["UPLOAD_FOLDER"] = chosen
app.logger.info("Using UPLOAD_FOLDER=%s", app.config["UPLOAD_FOLDER"])



upload_dir = _ensure_writable_dir(
    ([UPLOAD_ROOT_ENV] if UPLOAD_ROOT_ENV else []) + [INSTANCE_UPLOADS, "/tmp/policylens_uploads"]
)

#app.config['UPLOAD_FOLDER'] = upload_dir
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv("MAX_UPLOAD_MB", "16")) * 1024 * 1024  # 16MB default
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx'}
app.logger.info("Uploads directory set to: %s", app.config['UPLOAD_FOLDER'])

# -----------------------------------------------------------------------------
# Services
# -----------------------------------------------------------------------------
doc_processor = DocumentProcessor()
analysis_service = AnalysisService(doc_processor=doc_processor)

# -----------------------------------------------------------------------------
# Login
# -----------------------------------------------------------------------------
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# -----------------------------------------------------------------------------
# Logging (inject request_id so formatter never breaks)
# -----------------------------------------------------------------------------
class RequestIDFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        try:
            rid = getattr(g, 'request_id', None)
        except Exception:
            rid = None
        record.request_id = rid or '-'  # ensure field always exists
        return True

@app.before_request
def _assign_request_id():
    # one id per request; used in logs
    g.request_id = uuid.uuid4().hex[:8]

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
handler = logging.StreamHandler(sys.stdout)
handler.addFilter(RequestIDFilter())
formatter = logging.Formatter("%(asctime)s | %(levelname)s | %(name)s | req=%(request_id)s | %(message)s")
handler.setFormatter(formatter)
root = logging.getLogger()
root.setLevel(LOG_LEVEL)
root.handlers = [handler]

# Optional: startup sanity log (don’t rely on request_id here; filter handles it)
app.logger.info(
    "Startup: SECRET_KEY=%s, DB=%s",
    "set" if os.getenv("SECRET_KEY") else "missing",
    app.config['SQLALCHEMY_DATABASE_URI']
)



class RequestIdFilter(logging.Filter):
    def filter(self, record):
        # Always set a request_id so the formatter never fails
        record.request_id = "-"
        # Only read g if a request context exists
        if has_request_context():
            try:
                record.request_id = getattr(g, "request_id", "-")
            except Exception:
                # ultra-defensive: never let logging crash app startup
                record.request_id = "-"
        return True

root_logger = logging.getLogger()
root_logger.setLevel(LOG_LEVEL)

# Avoid duplicate handlers on reloads
if not any(isinstance(h, logging.StreamHandler) for h in root_logger.handlers):
    root_logger.addHandler(handler)

# Ensure our filter is on all handlers so %(request_id)s is always present
for h in root_logger.handlers:
    # prevent stacking multiple identical filters
    if not any(isinstance(f, RequestIdFilter) for f in getattr(h, "filters", [])):
        h.addFilter(RequestIdFilter())

app.logger.setLevel(LOG_LEVEL)
@app.get("/health")
def health():
    return {"ok": True}, 200


@app.before_request
def _assign_request_id_and_log():
    g.request_id = str(uuid.uuid4())[:8]
    app.logger.info(
        f"→ {request.method} {request.path} "
        f"args={dict(request.args)} json={(request.get_json(silent=True) or {})}"
    )

@app.after_request
def _log_response(resp):
    app.logger.info(f"← {request.method} {request.path} status={resp.status_code}")
    return resp

@app.teardown_request
def _teardown_request(exc):
    if exc:
        app.logger.exception(f"✖ Unhandled exception during request: {exc}")

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

with app.app_context():
    db.create_all()

# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
def can_view_all():
    """Check if current user can view all content (admin or view-only admin)"""
    return current_user.is_authenticated and (current_user.is_admin or current_user.is_view_only_admin)

def can_edit_all():
    """Check if current user can edit all content (admin but not view-only)"""
    return current_user.is_authenticated and current_user.is_admin and not current_user.is_view_only_admin

def media_allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

IMPORT_ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'doc', 'docx'}

def _canon_name(s: str) -> str:
    s = unicodedata.normalize('NFKD', s).encode('ascii','ignore').decode('ascii')
    s = s.strip().lower()
    s = s.replace('%20', ' ').replace('+', ' ').replace('_', ' ')
    name, dot, ext = s.rpartition('.')
    name = re.sub(r'[^a-z0-9\s-]', ' ', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return f"{name}{'.' if dot else ''}{ext}"

def _find_existing_file(filename: str) -> str | None:
    wanted = _canon_name(filename)
    search_dirs = [
        app.config.get('UPLOAD_FOLDER'),
        os.path.join(app.root_path, 'uploads', 'codebooks'),
        os.path.join(app.root_path, 'uploads'),
        os.path.join(app.root_path, 'uploaded_files'),
        os.path.join(app.root_path, 'frameworks'),
        os.path.join(app.root_path, 'policies'),
        os.path.join(app.root_path, 'frame_uploaded_files'),
        os.path.join(app.root_path, 'policy_uploaded_files'),
        os.path.join(app.root_path, 'data', 'frame_uploaded_files'),
        os.path.join(app.root_path, 'data', 'policy_uploaded_files'),
        os.path.join(app.root_path, 'static'),
    ]
    tried = []
    for d in filter(None, search_dirs):
        try:
            for f in os.listdir(d):
                if f == filename:
                    return os.path.join(d, f)
                if _canon_name(f) == wanted:
                    return os.path.join(d, f)
            tried.append(d)
        except FileNotFoundError:
            continue
    try:
        media = Media.query.filter_by(filename=filename).first()
        if not media:
            all_my_media = Media.query.filter((Media.user_id == current_user.id) | (Media.visibility == 'public')).all()
            for m in all_my_media:
                if _canon_name(m.filename) == wanted:
                    media = m
                    break
        if media:
            media_path = os.path.join(app.config['UPLOAD_FOLDER'], media.filename)
            if os.path.exists(media_path):
                return media_path
    except Exception:
        pass
    exts = {'.pdf', '.doc', '.docx'}
    for root, _, files in os.walk(app.root_path):
        if any(seg in root for seg in ('.git', '__pycache__', 'venv', '.venv', 'node_modules')):
            continue
        for f in files:
            _, ext = os.path.splitext(f)
            if ext.lower() not in exts:
                continue
            if _canon_name(f) == wanted:
                return os.path.join(root, f)
    app.logger.warning(f"File not found: '{filename}'. Looked in: {tried} (and DB + recursive scan)")
    return None

def _extract_vars_from_file(filepath: str) -> list[dict]:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        text = doc_processor.extract_text_from_pdf(filepath)
    elif ext in ('.doc', '.docx'):
        text = doc_processor.extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    if not (text or '').strip():
        raise ValueError("No text extracted from document")

    # vars_info = analysis_service.extract_codebook_codes(text, os.path.basename(filepath)) or []
    vars_info = analysis_service.extract_framework_variables(text, os.path.basename(filepath)) or []

    items = []
    for v in vars_info:
        name = (v.get('variable') or v.get('name') or '').strip()
        desc = (v.get('explanation') or v.get('description') or '').strip()
        if name:
            items.append({'name': name, 'description': desc})
    return items

class RequestCache:
    _instance = None
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance.cache = {}
        return cls._instance
    def add_request(self, request_id, ttl_seconds=30):
        self.cache[request_id] = datetime.utcnow() + timedelta(seconds=ttl_seconds)
        self.cleanup()
    def is_duplicate(self, request_id):
        self.cleanup()
        return request_id in self.cache
    def cleanup(self):
        now = datetime.utcnow()
        for rid in [r for r, t in self.cache.items() if t < now]:
            self.cache.pop(rid, None)

request_cache = RequestCache()

def generate_request_id(request_data) -> str:
    data_str = json.dumps(request_data, sort_keys=True)
    return hashlib.md5(data_str.encode('utf-8')).hexdigest()

# Codebook export helpers
# ---
def flatten_codes(codebook):
    """Flatten the code hierarchy into a list of records for exporting."""
    records = []
    for code in codebook.codes:
        records.append({
            'Level': 'Code',
            'Code': code.code,
            'Description': code.description or '',
            'Parent': '',
        })
        for subcode in code.subcodes:
            records.append({
                'Level': 'SubCode',
                'Code': f"  → {subcode.subcode}",
                'Description': subcode.description or '',
                'Parent': code.code,
            })
            for subsubcode in subcode.subsubcodes:
                records.append({
                    'Level': 'SubSubCode',
                    'Code': f"    → {subsubcode.subsubcode}",
                    'Description': subsubcode.description or '',
                    'Parent': subcode.subcode,
                })
    return records

def export_codebook_to_pdf(codebook):
    """Generate a PDF export of the codebook."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor('#003366'), spaceAfter=6)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=12, textColor=colors.HexColor('#003366'), spaceAfter=10)
    cell_style = ParagraphStyle('CellText', parent=styles['Normal'], fontSize=8, wordWrap='CJK')
    
    # Title and metadata
    elements.append(Paragraph(codebook.name, title_style))
    meta_text = f"<b>Created:</b> {codebook.created_at.strftime('%Y-%m-%d')} | <b>Codes:</b> {len(codebook.codes)}"
    elements.append(Paragraph(meta_text, styles['Normal']))
    if codebook.description:
        elements.append(Paragraph(f"<b>Description:</b> {codebook.description}", styles['Normal']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Codes table
    elements.append(Paragraph("Codes and Sub-Codes", heading_style))
    records = flatten_codes(codebook)
    
    # Build table data with wrapped text
    table_data = [['Level', 'Code', 'Description']]
    for rec in records:
        # Wrap descriptions in Paragraph for proper text wrapping
        desc_para = Paragraph(rec['Description'][:300], cell_style)  # Limit but let Paragraph handle wrapping
        table_data.append([rec['Level'], rec['Code'], desc_para])
    
    table = Table(table_data, colWidths=[0.9*inch, 1.8*inch, 2.8*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
    ]))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def export_codebook_to_docx(codebook):
    """Generate a DOCX export of the codebook."""
    doc = Document()
    
    # Title and metadata
    title = doc.add_heading(codebook.name, level=1)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    meta = doc.add_paragraph(f"Created: {codebook.created_at.strftime('%Y-%m-%d')} | Codes: {len(codebook.codes)}")
    meta.runs[0].font.size = Pt(10)
    
    if codebook.description:
        desc_para = doc.add_paragraph(codebook.description)
        desc_para.runs[0].italic = True
    
    doc.add_paragraph()  # Spacer
    
    # Codes section
    doc.add_heading("Codes and Sub-Codes", level=2)
    
    records = flatten_codes(codebook)
    
    # Add table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Level'
    header_cells[1].text = 'Code'
    header_cells[2].text = 'Description'
    
    # Style header
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add rows
    for rec in records:
        row_cells = table.add_row().cells
        row_cells[0].text = rec['Level']
        row_cells[1].text = rec['Code']
        row_cells[2].text = rec['Description'][:200]
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_codebook_to_excel(codebook):
    """Generate an Excel export of the codebook."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Codebook"
    
    # Header styling
    header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=12)
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    # Metadata section
    ws['A1'] = "Codebook Name"
    ws['B1'] = codebook.name
    ws['A2'] = "Created Date"
    ws['B2'] = codebook.created_at.strftime('%Y-%m-%d')
    ws['A3'] = "Total Codes"
    ws['B3'] = len(codebook.codes)
    if codebook.description:
        ws['A4'] = "Description"
        ws['B4'] = codebook.description
    
    # Bold metadata labels
    for cell in [ws['A1'], ws['A2'], ws['A3'], ws['A4']]:
        cell.font = Font(bold=True)
    
    # Codes table
    start_row = 6
    headers = ['Level', 'Code', 'Description', 'Parent Code']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal='center')
    
    records = flatten_codes(codebook)
    for row_num, rec in enumerate(records, start_row + 1):
        ws.cell(row=row_num, column=1).value = rec['Level']
        ws.cell(row=row_num, column=2).value = rec['Code']
        ws.cell(row=row_num, column=3).value = rec['Description']
        ws.cell(row=row_num, column=4).value = rec['Parent']
        
        for col in range(1, 5):
            ws.cell(row=row_num, column=col).border = border
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 40
    ws.column_dimensions['D'].width = 20
    
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# Auth
# -----------------------------------------------------------------------------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('dashboard'))
        flash('Invalid username or password', 'danger')
    return render_template('auth/login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        if password != confirm_password:
            flash('Passwords do not match', 'danger')
            return redirect(url_for('register'))
        if User.query.filter_by(username=username).first():
            flash('Username already exists', 'danger'); return redirect(url_for('register'))
        if User.query.filter_by(email=email).first():
            flash('Email already registered', 'danger'); return redirect(url_for('register'))
        user = User(username=username, email=email)
        user.set_password(password)
        db.session.add(user); db.session.commit()
        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))
    return render_template('auth/register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    """Request a password reset link via email."""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        if not email:
            flash('Please enter an email address.', 'danger')
            return redirect(url_for('forgot_password'))
        
        user = User.query.filter_by(email=email).first()
        if user:
            try:
                token = user.get_reset_token()
                reset_url = url_for('reset_password', token=token, _external=True)
                text_body = f"""To reset your password, visit the following link:
{reset_url}

This link will expire in 24 hours. If you did not request a password reset, please ignore this email.
"""
                html_body = render_template('auth/reset_email.html', user=user, reset_url=reset_url)
                msg = Message(
                    subject='Password Reset Request',
                    recipients=[user.email],
                    body=text_body,
                    html=html_body
                )
                mail.send(msg)
                flash('A password reset link has been sent to your email address.', 'info')
                return redirect(url_for('login'))
            except Exception as e:
                app.logger.error(f'Error sending password reset email: {str(e)}')
                flash('An error occurred. Please try again later.', 'danger')
        else:
            # Don't reveal if email exists (security)
            flash('If an account with that email exists, a reset link has been sent.', 'info')
            return redirect(url_for('login'))
    
    return render_template('auth/forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """Reset password using a valid token."""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    user = User.verify_reset_token(token)
    if not user:
        flash('The password reset link is invalid or has expired.', 'danger')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        if not password or not confirm_password:
            flash('Please fill in all fields.', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        if len(password) < 6:
            flash('Password must be at least 6 characters long.', 'danger')
            return redirect(url_for('reset_password', token=token))
        
        user.set_password(password)
        db.session.commit()
        flash('Your password has been reset. Please log in with your new password.', 'success')
        return redirect(url_for('login'))
    
    return render_template('auth/reset_password.html', token=token)

# Admin ability to reset user passwords
@app.route('/admin/users/<int:user_id>/reset-password', methods=['POST'])
@login_required
def admin_reset_user_password(user_id):
    """Admin can reset a user's password and send them a reset link."""
    if not current_user.is_admin:
        flash('You do not have permission to reset passwords.', 'danger')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    
    try:
        token = user.get_reset_token()
        reset_url = url_for('reset_password', token=token, _external=True)
        text_body = f"""An administrator has requested that you reset your password.
To reset your password, visit the following link:
{reset_url}

This link will expire in 24 hours. If you did not request this, please contact an administrator.
"""
        html_body = render_template('auth/reset_email.html', user=user, reset_url=reset_url, admin_reset=True)
        msg = Message(
            subject='Password Reset Request from Administrator',
            recipients=[user.email],
            body=text_body,
            html=html_body
        )
        mail.send(msg)
        flash(f'Password reset link sent to {user.email}.', 'success')
    except Exception as e:
        app.logger.error(f'Error sending password reset email: {str(e)}')
        flash('An error occurred sending the reset email.', 'danger')
    
    return redirect(request.referrer or url_for('dashboard'))

# -----------------------------------------------------------------------------
# Descriptors (media)
# -----------------------------------------------------------------------------
@app.route('/media/<int:media_id>/descriptors', methods=['GET'])
@login_required
def edit_descriptors(media_id):
    if current_user.is_view_only_admin:
        flash('View-only admins cannot edit descriptors.', 'danger')
        return redirect(url_for('list_media'))
    media = Media.query.get_or_404(media_id)
    if media.user_id != current_user.id:
        flash('You do not have permission to edit this media.', 'danger')
        return redirect(url_for('list_media'))
    return render_template('media/edit_descriptors.html', media=media)

@app.route('/media/<int:media_id>/descriptors/add', methods=['POST'])
@login_required
def add_descriptor(media_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot add descriptors'}), 403
    media = Media.query.get_or_404(media_id)
    if media.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Permission denied'}), 403
    data = request.get_json()
    key = data.get('key', '').strip()
    value = data.get('value', '').strip()
    if not key or not value:
        return jsonify({'success': False, 'message': 'Both key and value are required'}), 400
    existing = Descriptor.query.filter_by(media_id=media_id, key=key).first()
    if existing:
        return jsonify({'success': False, 'message': f'A descriptor with key "{key}" already exists for this media'}), 400
    descriptor = Descriptor(key=key, value=value, media_id=media_id)
    db.session.add(descriptor); db.session.commit()
    return jsonify({'success': True, 'descriptor': {'id': descriptor.id, 'key': descriptor.key, 'value': descriptor.value}})

@app.route('/media/<int:media_id>/descriptors/update', methods=['POST'])
@login_required
def update_descriptor(media_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot update descriptors'}), 403
    media = Media.query.get_or_404(media_id)
    data = request.get_json()
    if media.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Permission denied'}), 403
    descriptor_id = data.get('descriptor_id')
    key = data.get('key', '').strip()
    value = data.get('value', '').strip()
    if not descriptor_id or not key or not value:
        return jsonify({'success': False, 'message': 'Invalid request'}), 400
    descriptor = Descriptor.query.filter_by(id=descriptor_id, media_id=media_id).first()
    if not descriptor:
        return jsonify({'success': False, 'message': 'Descriptor not found'}), 404
    existing = Descriptor.query.filter(
        Descriptor.media_id == media_id, Descriptor.key == key, Descriptor.id != descriptor_id
    ).first()
    if existing:
        return jsonify({'success': False, 'message': f'A descriptor with key "{key}" already exists for this media'}), 400
    descriptor.key = key; descriptor.value = value
    db.session.commit()
    return jsonify({'success': True})

@app.route('/media/<int:media_id>/descriptors/delete', methods=['POST'])
@login_required
def delete_descriptor(media_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot delete descriptors'}), 403
    media = Media.query.get_or_404(media_id)
    data = request.get_json()
    if media.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Permission denied'}), 403
    descriptor_id = data.get('descriptor_id')
    if not descriptor_id:
        return jsonify({'success': False, 'message': 'Invalid request'}), 400
    descriptor = Descriptor.query.filter_by(id=descriptor_id, media_id=media_id).first()
    if descriptor:
        db.session.delete(descriptor); db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'message': 'Descriptor not found'}), 404

# -----------------------------------------------------------------------------
# Projects
# -----------------------------------------------------------------------------
@app.route('/projects')
@login_required
def list_projects():
    try:
        page = int(request.args.get('page', 1))
        if page < 1: page = 1
    except (ValueError, TypeError):
        page = 1
    search_query = request.args.get('q', '').strip()
    if can_view_all():
        query = Project.query
    else:
        query = Project.query.filter_by(owner_id=current_user.id)
    if search_query:
        search = f"%{search_query}%"
        query = query.filter(
            (Project.name.ilike(search)) |
            (Project.description.ilike(search) if Project.description is not None else False)
        )
    projects = query.order_by(Project.name).paginate(page=page, per_page=30, error_out=False)
    return render_template('projects/list.html', projects=projects, search_query=search_query)

@app.route('/projects/<int:project_id>')
@login_required
def view_project(project_id):
    project = Project.query.options(
        db.joinedload(Project.owner),
        db.joinedload(Project.media_items),
        db.joinedload(Project.codebooks),
        db.joinedload(Project.collaborators),
        db.joinedload(Project.excerpts).joinedload(Excerpt.media),
        db.joinedload(Project.excerpts).joinedload(Excerpt.user)
    ).get_or_404(project_id)
    if not can_view_all() and project.owner_id != current_user.id and not any(c.id == current_user.id for c in project.collaborators):
        flash('You do not have permission to view this project', 'danger')
        return redirect(url_for('list_projects'))
    media_items = project.media_items
    codebooks = project.codebooks
    collaborators = project.collaborators
    is_owner = project.owner_id == current_user.id
    if hasattr(project, 'excerpts'):
        project.excerpts.sort(key=lambda x: x.created_at, reverse=True)
    return render_template('projects/view.html',
                           project=project,
                           media_items=media_items,
                           codebooks=codebooks,
                           collaborators=collaborators,
                           is_owner=is_owner)

@app.route('/projects/new', methods=['GET', 'POST'])
@login_required
def create_project():
    if current_user.is_view_only_admin:
        flash('View-only admins cannot create projects.', 'danger')
        return redirect(url_for('list_projects'))
    form = ProjectForm()
    if form.validate_on_submit():
        project = Project(name=form.name.data, description=form.description.data, owner_id=current_user.id)
        db.session.add(project); db.session.commit()
        flash('Project created successfully!', 'success')
        return redirect(url_for('list_projects'))
    return render_template('projects/form.html', form=form, project=None)

@app.route('/projects/save', defaults={'project_id': None}, methods=['POST'])
@app.route('/projects/save/<int:project_id>', methods=['POST'])
@login_required
def save_project(project_id=None):
    if current_user.is_view_only_admin:
        flash('View-only admins cannot edit projects.', 'danger')
        return redirect(url_for('list_projects'))
    form = ProjectForm(request.form)
    if form.validate():
        try:
            if project_id:
                project = Project.query.get_or_404(project_id)
                if project.owner_id != current_user.id and not can_edit_all():
                    flash('You do not have permission to edit this project.', 'danger')
                    return redirect(url_for('list_projects'))
                project.name = form.name.data
                project.description = form.description.data
            else:
                project = Project(name=form.name.data, description=form.description.data, owner_id=current_user.id)
                db.session.add(project)
            db.session.commit()
            flash('Project saved successfully!', 'success')
            return redirect(url_for('list_projects'))
        except Exception as e:
            db.session.rollback()
            flash('An error occurred while saving the project.', 'danger')
    else:
        for field, errors in form.errors.items():
            for error in errors:
                flash(f"Error in {field}: {error}", 'danger')
    return render_template('projects/form.html', form=form, project=Project(id=project_id) if project_id else None)

@app.route('/projects/<int:project_id>/edit', methods=['GET'])
@login_required
def edit_project(project_id):
    if current_user.is_view_only_admin:
        flash('View-only admins cannot edit projects.', 'danger')
        return redirect(url_for('list_projects'))
    project = Project.query.get_or_404(project_id)
    if project.owner_id != current_user.id and not can_edit_all():
        flash('You do not have permission to edit this project.', 'danger')
        return redirect(url_for('list_projects'))
    form = ProjectForm(obj=project)
    return render_template('projects/form.html', form=form, project=project)

@app.route('/projects/<int:project_id>/add_collaborator', methods=['POST'])
@login_required
def add_collaborator(project_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot add collaborators'}), 403
    try:
        project = Project.query.get_or_404(project_id)
        if project.owner_id != current_user.id and not can_edit_all():
            return jsonify({'success': False, 'message': 'Only the project owner or an admin can add collaborators'}), 403
        data = request.get_json() or {}
        email = data.get('email')
        if not email:
            return jsonify({'success': False, 'message': 'Email is required'}), 400
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        if any(c.id == user.id for c in project.collaborators):
            return jsonify({'success': False, 'message': 'User is already a collaborator'}), 400
        if user.id == project.owner_id:
            return jsonify({'success': False, 'message': 'User is the project owner'}), 400
        project.collaborators.append(user)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Collaborator added successfully',
                        'collaborator': {'id': user.id, 'email': user.email,
                                         'name': user.name or user.email.split('@')[0],
                                         'created_at': user.created_at.isoformat() if user.created_at else None}})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error adding collaborator: {str(e)}', exc_info=True)
        return jsonify({'success': False, 'message': f'An error occurred while adding the collaborator: {str(e)}'}), 500

@app.route('/projects/<int:project_id>/remove_collaborator/<int:user_id>', methods=['POST'])
@login_required
def remove_collaborator(project_id, user_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot remove collaborators'}), 403
    project = Project.query.get_or_404(project_id)
    if project.owner_id != current_user.id and not can_edit_all():
        return jsonify({'success': False, 'message': 'Only the project owner or an admin can remove collaborators'}), 403
    collaborator = next((c for c in project.collaborators if c.id == user_id), None)
    if not collaborator:
        return jsonify({'success': False, 'message': 'User is not a collaborator'}), 404
    try:
        project.collaborators = [c for c in project.collaborators if c.id != user_id]
        db.session.commit()
        return jsonify({'success': True, 'message': 'Collaborator removed successfully', 'user_id': user_id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/projects/<int:project_id>/delete', methods=['POST'])
@login_required
def delete_project(project_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'message': 'View-only admins cannot delete projects'}), 403
    project = Project.query.get_or_404(project_id)
    if project.owner_id != current_user.id and not can_edit_all():
        return jsonify({'success': False, 'message': 'Permission denied'}), 403
    db.session.delete(project); db.session.commit()
    return jsonify({'success': True})

# -----------------------------------------------------------------------------
# Media
# -----------------------------------------------------------------------------
@app.route('/media')
@login_required
def list_media():
    # Get filter parameters
    project_id = request.args.get('project_id', type=int)
    search = request.args.get('search', '')
    
    # Get all projects for the sidebar
    projects = Project.query.filter(
        (Project.owner_id == current_user.id) | 
        (Project.collaborators.any(id=current_user.id))
    ).order_by(Project.name).all()
    
    # Query media files with optional project filter
    if can_view_all():
        query = Media.query
    else:
        query = Media.query.filter(
            (Media.user_id == current_user.id) | 
            (Media.visibility == 'public')
        )
    
    if project_id:
        query = query.filter(Media.project_id == project_id)
    
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            (Media.filename.ilike(search_term)) | 
            (Media.description.ilike(search_term))
        )
    
    # Get all media items (removed pagination for now)
    media_items = query.order_by(Media.uploaded_at.desc()).all()
    
    # Calculate total media count for "All Files" badge
    if can_view_all():
        all_media_count = Media.query.count()
    else:
        all_media_count = Media.query.filter(
            (Media.user_id == current_user.id) | 
            (Media.visibility == 'public')
        ).count()
    
    return render_template(
        'media/explorer.html',
        projects=projects,
        media_items=media_items,
        current_project_id=project_id,
        search_query=search,
        all_media_count=all_media_count
    )

@app.route('/media/upload', methods=['GET', 'POST'])
@login_required
def upload_media():
    if current_user.is_view_only_admin:
        flash('View-only admins cannot upload media.', 'danger')
        return redirect(url_for('list_media'))
    form = MediaUploadForm()
    form.project_id.choices = [(0, '-- No Project --')] + [(p.id, p.name) for p in Project.query.filter_by(owner_id=current_user.id).all()]
    if form.validate_on_submit():
        file = form.file.data
        if not file or file.filename == '':
            flash('No selected file', 'danger'); return redirect(request.url)
        if not media_allowed_file(file.filename):
            flash('File type not allowed', 'danger'); return redirect(request.url)

        # First check if file already exists in upload_dir by filename (quick check)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        filename = secure_filename(file.filename)
        if os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], filename)):
            flash(f'A file with this name already exists in the system - {filename}', 'warning')
            return redirect(request.url)

        # Read bytes to compute size and checksum for duplicate detection
        file_bytes = file.read()
        file_size = len(file_bytes)
        file.stream.seek(0)
        file_type = mimetypes.guess_type(file.filename)[0] or 'application/octet-stream'
        checksum = hashlib.sha256(file_bytes).hexdigest()

        # Look for existing media with same size and type; compare checksums on disk
        candidates = Media.query.filter_by(file_size=file_size, file_type=file_type).all()
        for cand in candidates:
            try:
                existing_path = os.path.join(app.config['UPLOAD_FOLDER'], cand.filename)
                if os.path.exists(existing_path):
                    with open(existing_path, 'rb') as fh:
                        existing_hash = hashlib.sha256(fh.read()).hexdigest()
                    if existing_hash == checksum:
                        # Optionally compare descriptors if provided to be extra-sure
                        provided_desc = {request.form.get(k): request.form.get(f"descriptor_value_{k.split('_')[-1]}")
                                         for k in request.form.keys() if k.startswith('descriptor_key_')}
                        if provided_desc:
                            existing_desc = {d.key: d.value for d in cand.descriptors}
                            if not all(existing_desc.get(k) == v for k, v in provided_desc.items()):
                                continue
                        flash('A file identical to this one already exists in the system.', 'warning')
                        return redirect(url_for('view_media', media_id=cand.id))
            except Exception:
                continue

        # Save file to upload folder
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        media = Media(
            filename=filename,
            description=form.description.data,
            category=form.category.data,
            type=form.type.data,
            visibility=form.visibility.data,
            file_type=file_type,
            file_size=file_size,
            user_id=current_user.id,
            project_id=form.project_id.data if form.project_id.data != 0 else None
        )
        for key in request.form.keys():
            if key.startswith('descriptor_key_'):
                idx = key.replace('descriptor_key_', '')
                desc_key = request.form.get(f'descriptor_key_{idx}')
                desc_value = request.form.get(f'descriptor_value_{idx}')
                if desc_key and desc_value:
                    db.session.add(Descriptor(key=desc_key, value=desc_value, media=media))
        db.session.add(media)
        if form.project_id.data and form.project_id.data != 0:
            project = Project.query.get(form.project_id.data)
            if project:
                project.media_count = project.media_count + 1 if project.media_count else 1
        db.session.commit()
        flash('File uploaded successfully!', 'success')
        return redirect(url_for('view_media', media_id=media.id))
    return render_template('media/upload.html', form=form, allowed_extensions=app.config['ALLOWED_EXTENSIONS'])

@app.route('/media/<int:media_id>')
@login_required
def view_media(media_id):
    media = Media.query.get_or_404(media_id)
    if not can_view_all() and media.user_id != current_user.id and media.visibility != 'public':
        flash('You do not have permission to view this media.', 'danger')
        return redirect(url_for('list_media'))
    descriptors = {d.key: d.value for d in media.descriptors}
    return render_template('media/view.html', media=media, descriptors=descriptors)

@app.route('/media/<int:media_id>/download')
@login_required
def download_media(media_id):
    media = Media.query.get_or_404(media_id)
    if not can_view_all() and media.user_id != current_user.id and media.visibility != 'public':
        flash('You do not have permission to download this media.', 'danger')
        return redirect(url_for('list_media'))
    return send_from_directory(app.config['UPLOAD_FOLDER'], media.filename, as_attachment=True)

@app.route('/media/<int:media_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_media(media_id):
    media = Media.query.get_or_404(media_id)
    if current_user.is_view_only_admin:
        flash('View-only admins cannot edit media.', 'danger')
        return redirect(url_for('list_media'))
    if not current_user.is_admin and media.user_id != current_user.id:
        flash('You do not have permission to edit this media.', 'danger')
        return redirect(url_for('list_media'))
    
    if request.method == 'POST':
        media.filename = request.form.get('filename', media.filename)
        media.description = request.form.get('description', media.description)
        media.category = request.form.get('category', media.category)
        media.visibility = request.form.get('visibility', media.visibility)
        
        # Update descriptors
        # First, remove existing descriptors to avoid duplicates/orphans if logic is simple replacement
        # Or better, update existing and add new. For simplicity in this context, let's clear and re-add 
        # or just handle keys. 
        # Let's use a simpler approach: Delete all for this media and re-add from form.
        Descriptor.query.filter_by(media_id=media.id).delete()
        
        for key in request.form.keys():
            if key.startswith('descriptor_key_'):
                idx = key.replace('descriptor_key_', '')
                desc_key = request.form.get(f'descriptor_key_{idx}')
                desc_value = request.form.get(f'descriptor_value_{idx}')
                if desc_key and desc_value:
                    db.session.add(Descriptor(key=desc_key, value=desc_value, media=media))
        
        db.session.commit()
        flash('Media updated successfully!', 'success')
        return redirect(url_for('view_media', media_id=media.id))
        
    return render_template('media/edit.html', media=media)

@app.route('/media/<int:media_id>/delete', methods=['POST'])
@login_required
def delete_media(media_id):
    if current_user.is_view_only_admin:
        return jsonify({'success': False, 'error': 'View-only admins cannot delete media'}), 403
    media = Media.query.get_or_404(media_id)
    if media.user_id != current_user.id:
        return jsonify({'success': False, 'error': 'Permission denied'}), 403
    project_id = media.project_id
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], media.filename)
    if os.path.exists(file_path):
        try: os.remove(file_path)
        except Exception as e: app.logger.error(f'Error deleting file {file_path}: {e}')
    db.session.delete(media)
    if project_id:
        project = Project.query.get(project_id)
        if project and project.media_count > 0:
            project.media_count -= 1
    db.session.commit()
    flash('Media deleted successfully!', 'success')
    return redirect(url_for('list_media'))

@app.route('/uploads/<filename>')
@login_required
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# -----------------------------------------------------------------------------
# Dashboard / Profile
# -----------------------------------------------------------------------------
@app.route('/dashboard')
@login_required
def dashboard():
    recent_media = Media.query.filter_by(user_id=current_user.id).order_by(Media.uploaded_at.desc()).limit(5).all()
    project_count = Project.query.filter_by(owner_id=current_user.id).count()
    total_storage = db.session.query(db.func.sum(Media.file_size)).filter(Media.user_id == current_user.id).scalar() or 0
    return render_template('dashboard.html', recent_media=recent_media, project_count=project_count, total_storage=total_storage)

@app.route('/research')
@login_required
def research():
    return render_template('research/index.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    bio_form = ProfileUpdateForm(meta={'csrf': False})
    password_form = PasswordUpdateForm(meta={'csrf': False})
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        if form_type == 'bio':
            bio_form = ProfileUpdateForm(request.form, meta={'csrf': False})
            if bio_form.validate():
                try:
                    current_user.name = bio_form.name.data
                    current_user.age_group = bio_form.age_group.data or None
                    current_user.sex = bio_form.sex.data or None
                    current_user.industry = bio_form.industry.data or None
                    current_user.organization = bio_form.organization.data or None
                    db.session.commit()
                    flash('Your profile has been updated!', 'success')
                    return redirect(url_for('profile'))
                except Exception as e:
                    db.session.rollback()
                    flash(f'Error updating profile: {str(e)}', 'danger')
            else:
                for field, errors in bio_form.errors.items():
                    for error in errors:
                        flash(f"{getattr(bio_form, field).label.text}: {error}", 'danger')
        elif form_type == 'password':
            password_form = PasswordUpdateForm(request.form, meta={'csrf': False})
            if password_form.validate():
                if not check_password_hash(current_user.password_hash, password_form.current_password.data):
                    flash('Current password is incorrect', 'danger')
                else:
                    current_user.set_password(password_form.new_password.data)
                    db.session.commit()
                    flash('Your password has been updated!', 'success')
                return redirect(url_for('profile') + '#security')
            else:
                for field, errors in password_form.errors.items():
                    for error in errors:
                        flash(f"{getattr(password_form, field).label.text}: {error}", 'danger')
    if request.method == 'GET':
        bio_form.name.data = current_user.name or ''
        bio_form.email.data = current_user.email or ''
        bio_form.age_group.data = current_user.age_group or ''
        bio_form.sex.data = current_user.sex or ''
        bio_form.industry.data = current_user.industry or ''
        bio_form.organization.data = current_user.organization or ''
    return render_template('profile.html', bio_form=bio_form, password_form=password_form)

# -----------------------------------------------------------------------------
# Codebooks: AJAX endpoints
# -----------------------------------------------------------------------------
@app.route('/api/codebook/<int:codebook_id>/update_item', methods=['POST'])
@login_required
def update_codebook_item(codebook_id):
    data = request.get_json() or {}
    item_type = data.get('type'); item_id = data.get('id')
    name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip()
    if not item_type or not item_id or not name:
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400
    codebook = Codebook.query.get_or_404(codebook_id)
    if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        if item_type == 'code':
            code = Code.query.filter_by(id=item_id, codebook_id=codebook_id).first_or_404()
            code.code, code.description = name, description
        elif item_type == 'subcode':
            subcode = SubCode.query.join(Code).filter(
                SubCode.id == item_id, Code.codebook_id == codebook_id
            ).first_or_404()
            subcode.subcode, subcode.description = name, description
        elif item_type == 'subsubcode':
            subsubcode = SubSubCode.query.join(SubCode, Code).filter(
                SubSubCode.id == item_id, Code.codebook_id == codebook_id
            ).first_or_404()
            subsubcode.subsubcode, subsubcode.description = name, description
        else:
            return jsonify({'success': False, 'message': 'Invalid item type'}), 400
        db.session.commit()
        return jsonify({'success': True, 'message': f'{item_type.capitalize()} updated successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Error updating {item_type}: {str(e)}'}), 500

@app.route('/api/codebook/<int:codebook_id>/save_item', methods=['POST'])
@login_required
def save_codebook_item(codebook_id):
    if not request.is_json:
        return jsonify({'success': False, 'message': 'Invalid request format'}), 400
    data = request.get_json()
    request_id = request.headers.get('X-Request-ID') or generate_request_id(data)
    if request_cache.is_duplicate(request_id):
        return jsonify({'success': False, 'message': 'Duplicate request detected'}), 409
    request_cache.add_request(request_id)
    level = data.get('level'); name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip(); parent_id = data.get('parent_id')
    if not name:
        return jsonify({'success': False, 'message': 'Name is required'}), 400
    if level in ['subcode', 'subsubcode'] and not parent_id:
        return jsonify({'success': False, 'message': 'Parent ID is required'}), 400
    try:
        codebook = Codebook.query.get_or_404(codebook_id)
        if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        if level == 'code':
            if Code.query.filter_by(code=name, codebook_id=codebook_id).first():
                return jsonify({'success': False, 'message': 'A code with this name already exists'}), 400
            new_item = Code(code=name, description=description, codebook_id=codebook_id)
        elif level == 'subcode':
            code = Code.query.get_or_404(parent_id)
            if code.codebook_id != codebook_id:
                return jsonify({'success': False, 'message': 'Invalid parent code'}), 400
            if SubCode.query.filter_by(subcode=name, code_id=parent_id).first():
                return jsonify({'success': False, 'message': 'A subcode with this name already exists under the selected code'}), 400
            new_item = SubCode(subcode=name, description=description, code_id=parent_id)
        elif level == 'subsubcode':
            subcode = SubCode.query.get_or_404(parent_id)
            if subcode.code.codebook_id != codebook_id:
                return jsonify({'success': False, 'message': 'Invalid parent subcode'}), 400
            if SubSubCode.query.filter_by(subsubcode=name, subcode_id=parent_id).first():
                return jsonify({'success': False, 'message': 'A sub-subcode with this name already exists under the selected subcode'}), 400
            new_item = SubSubCode(subsubcode=name, description=description, subcode_id=parent_id)
        else:
            return jsonify({'success': False, 'message': 'Invalid level'}), 400
        db.session.add(new_item); db.session.commit()
        return jsonify({'success': True, 'message': f'{level.capitalize()} added successfully',
                        'item': {'id': new_item.id, 'name': name, 'description': description,
                                 'level': level, 'parent_id': parent_id if level != 'code' else None}})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error saving {level}: {str(e)}')
        return jsonify({'success': False, 'message': f'Failed to save {level}: {str(e)}'}), 500


@app.route('/api/codebook/<int:codebook_id>/delete_item', methods=['POST'])
@login_required
def delete_codebook_item(codebook_id):
    """Delete a code, subcode, or subsubcode."""
    try:
        codebook = db.session.get(Codebook, codebook_id)
        if not codebook or (codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403

        data = request.get_json() or {}
        item_type = data.get('type', '').lower()
        try:
            item_id = int(data.get('id', 0))
        except (ValueError, TypeError):
            item_id = None

        if not item_type or not item_id:
            return jsonify({'success': False, 'message': 'Missing type or id'}), 400

        if item_type == 'code':
            item = Code.query.get_or_404(item_id)
            if item.codebook_id != codebook_id:
                return jsonify({'success': False, 'message': 'Item does not belong to this codebook'}), 403
            db.session.delete(item)
        elif item_type == 'subcode':
            item = SubCode.query.get_or_404(item_id)
            code = Code.query.get(item.code_id)
            if not code or code.codebook_id != codebook_id:
                return jsonify({'success': False, 'message': 'Item does not belong to this codebook'}), 403
            db.session.delete(item)
        elif item_type == 'subsubcode':
            item = SubSubCode.query.get_or_404(item_id)
            subcode = SubCode.query.get(item.subcode_id)
            code = Code.query.get(subcode.code_id) if subcode else None
            if not code or code.codebook_id != codebook_id:
                return jsonify({'success': False, 'message': 'Item does not belong to this codebook'}), 403
            db.session.delete(item)
        else:
            return jsonify({'success': False, 'message': f'Unknown type: {item_type}'}), 400

        db.session.commit()
        return jsonify({'success': True, 'message': f'{item_type.capitalize()} deleted successfully'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error deleting {item_type}: {str(e)}')
        return jsonify({'success': False, 'message': f'Failed to delete item: {str(e)}'}), 500


# -----------------------------------------------------------------------------
# Codebook import (LLM-based)
# -----------------------------------------------------------------------------
@app.route('/api/codebook/import', methods=['GET', 'POST'])
@login_required
def import_codebook():
    """
    Import variables from a document into a codebook.

    POST: upload a new file (csv/xls/xlsx/pdf/doc/docx), extract variables.
          If a codebook_id is provided (in form-data or query string), save items
          into that codebook and return JSON. Otherwise, return JSON with a preview.

    GET:  ?file=<existing filename>&codebook_id=<id>
          Extract from an existing PDF/DOC/DOCX already on disk.
          If codebook_id is provided, save + flash + redirect to edit page.
          Otherwise, return JSON preview (no DB writes).
    """
    try:
        # Resolve codebook_id (POST may send it in multipart form; GET uses query string)
        if request.method == 'POST':
            codebook_id = request.form.get('codebook_id', type=int) or request.args.get('codebook_id', type=int)
        else:
            codebook_id = request.args.get('codebook_id', type=int)

        codebook = db.session.get(Codebook, codebook_id) if codebook_id else None
        if codebook and (codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403

        # ---------- GET: extract from an existing on-disk file (PDF/DOC/DOCX) ----------
        if request.method == 'GET':
            fname = (request.args.get('file') or '').strip()
            if not fname:
                return jsonify({'success': False, 'message': 'Missing file query parameter'}), 400

            filepath = _find_existing_file(fname)
            if not filepath:
                return jsonify({'success': False, 'message': f'File not found: {fname}'}), 404

            items = _extract_vars_from_file(filepath)  # LLM extraction from PDF/DOCX

            if codebook:
                added = 0
                for it in items:
                    if not Code.query.filter_by(codebook_id=codebook.id, code=it['name']).first():
                        db.session.add(Code(code=it['name'], description=it['description'], codebook_id=codebook.id))
                        added += 1
                db.session.commit()
                flash(f'Imported {added} codes from {fname}', 'success')
                return redirect(url_for('edit_codebook_codes', codebook_id=codebook.id))

            return jsonify({'success': True, 'filename': fname, 'items': items})

        # ---------- POST: handle upload then extract (CSV/XLS/XLSX/PDF/DOC/DOCX) ----------
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file part'}), 400
        file = request.files['file']
        if not file or not file.filename:
            return jsonify({'success': False, 'message': 'No selected file'}), 400

        ext = file.filename.rsplit('.', 1)[-1].lower()
        if ext not in IMPORT_ALLOWED_EXTENSIONS:
            return jsonify({'success': False, 'message': f'Invalid file type: .{ext}'}), 400

        # ✅ save under the same writable root as other uploads
        upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'codebooks')
        os.makedirs(upload_dir, exist_ok=True)

        # First check if file already exists by filename (quick check)
        base, ext_dot = os.path.splitext(secure_filename(file.filename))
        filename = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext_dot}"
        filepath = os.path.join(upload_dir, filename)
        if os.path.exists(filepath):
            if codebook:
                flash(f'File already exists: {filename}', 'warning')
                return redirect(url_for('edit_codebook_codes', codebook_id=codebook.id))
            else:
                return jsonify({'success': False, 'message': f'File already exists: {filename}'}), 409

        file.save(filepath)

        items = []
        ext_lower = ext_dot.lower()

        # Spreadsheets parsed directly; others go through LLM-based extractor
        if ext_lower in ('.csv', '.xls', '.xlsx'):
            try:
                df = pd.read_csv(filepath) if ext_lower == '.csv' else pd.read_excel(filepath)
                df.columns = [str(c).strip().lower() for c in df.columns]
                name_col = next((c for c in df.columns if c in ('name', 'code', 'variable')), None)
                desc_col = next((c for c in df.columns if c in ('description', 'desc', 'explanation')), None)
                if not name_col:
                    return jsonify({'success': False, 'message': 'Spreadsheet must contain a Name/Code/Variable column'}), 400
                for _, row in df.iterrows():
                    n = (str(row.get(name_col, '')).strip())
                    d = (str(row.get(desc_col, '')).strip()) if desc_col else ''
                    if n:
                        items.append({'name': n, 'description': d})
            except Exception as e:
                return jsonify({'success': False, 'message': f'Failed to parse spreadsheet: {e}'}), 400
        else:
            items = _extract_vars_from_file(filepath)  # PDF/DOC/DOCX via LLM

        # If a target codebook is provided on POST, save and return JSON for the modal flow
        if codebook:
            added = 0
            for it in items:
                if not Code.query.filter_by(codebook_id=codebook.id, code=it['name']).first():
                    db.session.add(Code(code=it['name'], description=it['description'], codebook_id=codebook.id))
                    added += 1
            db.session.commit()
            # JSON (no redirect) so the frontend can display status and reload
            return jsonify({'success': True, 'message': f'Imported {added} codes from {filename}', 'imported': added})

        # No codebook_id: return a preview (unchanged behavior)
        return jsonify({'success': True, 'filename': filename, 'items': items})

    except Exception as e:
        app.logger.error(f'Error processing import: {str(e)}', exc_info=True)
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

# (rest of your file continues unchanged)

# -----------------------------------------------------------------------------
# Codebook pages
# -----------------------------------------------------------------------------
@app.route('/codebooks')
@login_required
def list_codebooks():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '').strip()
    query = Codebook.query.options(joinedload(Codebook.project))
    # Admins see all codebooks; regular users see only their own
    if not current_user.is_admin:
        query = query.filter_by(user_id=current_user.id)
    if search:
        s = f"%{search}%"
        query = query.filter((Codebook.name.ilike(s)) | (Codebook.description.ilike(s)))
    codebooks = query.order_by(Codebook.name).paginate(page=page, per_page=30, error_out=False)
    return render_template('codebooks/list.html', codebooks=codebooks, search=search)

@app.route('/codebooks/new', methods=['GET', 'POST'])
@login_required
def create_codebook():
    from models.models import Project  # local import to avoid circulars in some setups
    form = CodebookForm()
    projects = Project.query.filter_by(owner_id=current_user.id).all()
    form.project_id.choices = [('', 'Select a project (optional)')] + [(str(p.id), p.name) for p in projects]
    if form.validate_on_submit():
        try:
            codebook = Codebook(
                name=form.name.data,
                description=form.description.data,
                user_id=current_user.id,
                project_id=form.project_id.data if form.project_id.data else None
            )
            db.session.add(codebook); db.session.commit()
            flash('Codebook created successfully!', 'success')
            return redirect(url_for('list_codebooks'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating codebook: {str(e)}', 'danger')
    return render_template('codebooks/form.html', form=form, title='Create Codebook')

@app.route('/codebooks/<int:codebook_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_codebook(codebook_id):
    from models.models import Project
    codebook = Codebook.query.get_or_404(codebook_id)
    if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True: abort(403)
    form = CodebookForm(obj=codebook)
    projects = Project.query.filter_by(owner_id=current_user.id).all()
    form.project_id.choices = [('', 'Select a project (optional)')] + [(str(p.id), p.name) for p in projects]
    if codebook.project_id:
        form.project_id.data = str(codebook.project_id)
    if form.validate_on_submit():
        try:
            codebook.name = form.name.data
            codebook.description = form.description.data
            codebook.project_id = form.project_id.data if form.project_id.data else None
            db.session.commit()
            flash('Codebook updated successfully!', 'success')
            return redirect(url_for('list_codebooks'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating codebook: {str(e)}', 'danger')
    return render_template('codebooks/form.html', form=form, title='Edit Codebook', codebook=codebook)

@app.route('/codebooks/<int:codebook_id>/delete', methods=['POST'])
@login_required
def delete_codebook(codebook_id):
    codebook = Codebook.query.get_or_404(codebook_id)
    if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True: abort(403)
    try:
        db.session.delete(codebook); db.session.commit()
        flash('Codebook deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting codebook: {str(e)}', 'danger')
    return redirect(url_for('list_codebooks'))

@app.route('/codebooks/<int:codebook_id>/export/<format>', methods=['GET'])
@login_required
def export_codebook(codebook_id, format):
    """Export codebook in PDF, DOCX, or XLSX format."""
    codebook = Codebook.query.options(
        joinedload(Codebook.codes).joinedload(Code.subcodes).joinedload(SubCode.subsubcodes)
    ).get_or_404(codebook_id)
    
    if codebook.user_id != current_user.id and not current_user.is_admin:
        abort(403)
    
    format = format.lower()
    if format == 'pdf':
        buffer = export_codebook_to_pdf(codebook)
        filename = f"{codebook.name.replace(' ', '_')}.pdf"
        return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)
    elif format == 'docx':
        buffer = export_codebook_to_docx(codebook)
        filename = f"{codebook.name.replace(' ', '_')}.docx"
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                        as_attachment=True, download_name=filename)
    elif format == 'xlsx':
        buffer = export_codebook_to_excel(codebook)
        filename = f"{codebook.name.replace(' ', '_')}.xlsx"
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        as_attachment=True, download_name=filename)
    else:
        flash('Invalid export format.', 'danger')
        return redirect(url_for('view_codebook', codebook_id=codebook_id))

@app.route('/code/<int:code_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_code(code_id):
    code = Code.query.get_or_404(code_id)
    codebook = code.codebook
    if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True: abort(403)
    if request.method == 'POST':
        try:
            code.code = request.form.get('code', '').strip()
            code.description = request.form.get('description', '').strip()
            db.session.commit()
            flash('Code updated successfully!', 'success')
            return redirect(url_for('edit_codebook_codes', codebook_id=codebook.id))
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating code: {str(e)}', 'danger')
    return render_template('codebooks/edit_code.html', code=code, codebook=codebook)

@app.route('/code/<int:code_id>/subcode', methods=['POST'])
@login_required
def add_subcode(code_id):
    code = Code.query.get_or_404(code_id)
    if code.codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    data = request.get_json() or {}
    if 'name' not in data:
        return jsonify({'success': False, 'message': 'Subcode name is required'}), 400
    try:
        subcode = SubCode(subcode=data['name'], description=data.get('description', ''), code_id=code_id)
        db.session.add(subcode); db.session.commit()
        return jsonify({'success': True, 'message': 'Subcode added successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/code/<int:code_id>/delete', methods=['POST'])
@login_required
def delete_code(code_id):
    code = Code.query.get_or_404(code_id)
    if code.codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True: abort(403)
    try:
        db.session.delete(code); db.session.commit()
        return '', 204
    except Exception as e:
        db.session.rollback()
        return str(e), 500

@app.route('/codebooks/<int:codebook_id>/edit_codes', methods=['GET', 'POST'])
@login_required
def edit_codebook_codes(codebook_id):
    # If a file param is provided, auto-import (LLM) and then redirect back here
    incoming = (request.args.get('file') or '').strip()
    if incoming:
        return redirect(url_for('import_codebook', codebook_id=codebook_id, file=incoming))

    codebook = Codebook.query.options(
        joinedload(Codebook.codes).joinedload(Code.subcodes).joinedload(SubCode.subsubcodes)
    ).get_or_404(codebook_id)
    if codebook.user_id != current_user.id and not current_user.is_admin and current_user.is_view_only_admin == True: abort(403)

    if request.method == 'POST' and request.is_json:
        try:
            data = request.get_json() or {}
            level = data.get('level'); name = (data.get('name') or '').strip()
            description = (data.get('description') or '').strip(); parent_id = data.get('parent_id')
            if not name: return jsonify({'success': False, 'message': 'Name is required'}), 400
            if level == 'code':
                new_code = Code(code=name, description=description, codebook_id=codebook_id)
                db.session.add(new_code); db.session.flush()
                return jsonify({'success': True, 'message': 'Code added successfully',
                                'item': {'id': new_code.id, 'name': new_code.code,
                                         'description': new_code.description, 'level': 'code'}})
            elif level == 'subcode':
                code = Code.query.get_or_404(parent_id)
                if code.codebook_id != codebook_id:
                    return jsonify({'success': False, 'message': 'Invalid parent code'}), 400
                new_subcode = SubCode(subcode=name, description=description, code_id=parent_id)
                db.session.add(new_subcode); db.session.flush()
                return jsonify({'success': True, 'message': 'Subcode added successfully',
                                'item': {'id': new_subcode.id, 'name': new_subcode.subcode,
                                         'description': new_subcode.description, 'level': 'subcode',
                                         'parent_id': parent_id}})
            elif level == 'subsubcode':
                subcode = SubCode.query.get_or_404(parent_id)
                if subcode.code.codebook_id != codebook_id:
                    return jsonify({'success': False, 'message': 'Invalid parent subcode'}), 400
                new_subsubcode = SubSubCode(subsubcode=name, description=description, subcode_id=parent_id)
                db.session.add(new_subsubcode); db.session.flush()
                return jsonify({'success': True, 'message': 'Sub-subcode added successfully',
                                'item': {'id': new_subsubcode.id, 'name': new_subsubcode.subsubcode,
                                         'description': new_subsubcode.description, 'level': 'subsubcode',
                                         'parent_id': parent_id}})
            else:
                return jsonify({'success': False, 'message': 'Invalid level'}), 400
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': str(e)}), 500

    return render_template('codebooks/edit_codes.html', codebook=codebook)

# -----------------------------------------------------------------------------
# Excerpts
# -----------------------------------------------------------------------------
@app.route('/api/excerpts/<int:excerpt_id>', methods=['GET'])
@login_required
def get_excerpt(excerpt_id):
    try:
        result = db.session.query(
            Excerpt,
            Media.filename.label('media_name'),
            User.username.label('user_name')
        ).outerjoin(Media, Excerpt.media_id == Media.id)\
         .outerjoin(User, Excerpt.user_id == User.id)\
         .filter(Excerpt.id == excerpt_id).first()
        if not result or not result.Excerpt:
            return jsonify({'success': False, 'message': 'Excerpt not found'}), 404
        project = Project.query.get(result.Excerpt.project_id)
        if not project or (not can_view_all() and project.owner_id != current_user.id and not any(c.id == current_user.id for c in project.collaborators)):
            return jsonify({'success': False, 'message': 'You do not have permission to view this excerpt'}), 403
        data = {
            'id': result.Excerpt.id,
            'project_id': result.Excerpt.project_id,
            'media_id': result.Excerpt.media_id,
            'media_name': result.media_name,
            'codebook_id': result.Excerpt.codebook_id,
            'code': result.Excerpt.code,
            'subcode': result.Excerpt.subcode,
            'excerpt': result.Excerpt.excerpt,
            'explanation': result.Excerpt.explanation,
            'user_id': result.Excerpt.user_id,
            'user_name': result.user_name,
            'created_at': result.Excerpt.created_at.isoformat() if result.Excerpt.created_at else None
        }
        return jsonify({'success': True, 'excerpt': data})
    except Exception as e:
        app.logger.error(f'Error fetching excerpt: {str(e)}')
        return jsonify({'success': False, 'message': 'An error occurred while fetching the excerpt'}), 500

# app.py (add after existing routes)
# -----------------------------------------------------------------------------
# Excerpts - scoped listing
# -----------------------------------------------------------------------------
@app.route('/get_excerpts', methods=['POST'])
@login_required
def get_excerpts():
    try:
        data = request.get_json() or {}
        project_id = data.get('project_id')
        scope = (data.get('scope') or 'mine').lower()  # UI may send this; we still filter 'mine' below

        app.logger.info(f"[get_excerpts] user={current_user.id} project={project_id} scope_req={scope}")

        if not project_id:
            app.logger.warning("[get_excerpts] missing project_id")
            return jsonify({'error': 'Missing required field: project_id'}), 400

        project = db.session.get(Project, project_id)
        if not project:
            app.logger.warning(f"[get_excerpts] project not found: {project_id}")
            return jsonify({'error': 'Project not found'}), 404

        if not can_view_all() and project.owner_id != current_user.id and not any(c.id == current_user.id for c in project.collaborators):
            app.logger.warning(f"[get_excerpts] unauthorized user={current_user.id} project={project_id}")
            return jsonify({'error': 'Unauthorized to view excerpts for this project'}), 403

        # Build query based on scope
        q = (
            db.session.query(Excerpt, Media.filename, Media.id.label('media_id'))
            .outerjoin(Media, Excerpt.media_id == Media.id)
            .filter(Excerpt.project_id == project_id)
        )
        
        # Filter by scope unless user is admin
        if not can_view_all():
            if scope == 'mine':
                q = q.filter(Excerpt.user_id == current_user.id)
            elif scope == 'team':
                # Team: owner + collaborators
                team_ids = [project.owner_id] + [c.id for c in project.collaborators]
                q = q.filter(Excerpt.user_id.in_(team_ids))
            # 'all' scope shows everything (no additional filter)
        
        q = q.order_by(Excerpt.created_at.desc())
        rows = q.all()
        app.logger.info(f"[get_excerpts] fetched={len(rows)}")

        excerpts_data = [{
            'id': ex.id,
            'media_id': media_id,
            'media_filename': media_filename or 'N/A',
            'code': ex.code or 'N/A',
            'subcode': ex.subcode or 'N/A',
            'excerpt': ex.excerpt or 'N/A',
            'explanation': ex.explanation or 'N/A',
            'created_at': ex.created_at.isoformat() if ex.created_at else None
        } for ex, media_filename, media_id in rows]

        return jsonify({'success': True, 'excerpts': excerpts_data})
    except Exception as e:
        app.logger.exception(f"[get_excerpts] failure: {e}")
        return jsonify({'error': f'Failed to fetch excerpts: {str(e)}'}), 500


@csrf.exempt
@app.route('/save_excerpt', methods=['POST'])
@login_required
def save_excerpt():
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No data provided'}), 400
    try:
        data = request.get_json() or {}
        project_id = data.get('project_id')
        excerpts = data.get('excerpts', [])

        app.logger.info(f"[save_excerpt] user={current_user.id} project={project_id} batch={len(excerpts)}")

        if not project_id or not excerpts:
            app.logger.warning("[save_excerpt] missing project_id or excerpts")
            return jsonify({'message': 'Missing project_id or excerpts', 'success': False}), 400

        project = db.session.get(Project, project_id)
        if not project:
            app.logger.warning(f"[save_excerpt] project not found: {project_id}")
            return jsonify({'message': 'Project not found', 'success': False}), 404
        if project.owner_id != current_user.id and current_user not in project.collaborators:
            app.logger.warning(f"[save_excerpt] unauthorized user={current_user.id} project={project_id}")
            return jsonify({'message': 'Unauthorized to save excerpts for this project', 'success': False}), 403

        excerpt_ids = []
        for excerpt in excerpts:
            media_id = excerpt.get('media_id')
            codebook_id = excerpt.get('codebook_id')
            code = excerpt.get('code')
            subcode = excerpt.get('subcode')
            excerpt_text = excerpt.get('excerpt')  # Corrected key
            explanation = excerpt.get('explanation')

            if not all([media_id, codebook_id, code, excerpt_text]):
                app.logger.warning(f"[save_excerpt] invalid excerpt: media_id={media_id} codebook_id={codebook_id} code={code}")
                continue

            media = db.session.get(Media, media_id)
            codebook = db.session.get(Codebook, codebook_id)
            if not media or not codebook:
                app.logger.warning(f"[save_excerpt] invalid media_id={media_id} or codebook_id={codebook_id}")
                continue
            if media.project_id != project_id or codebook.project_id != project_id:
                app.logger.warning(f"[save_excerpt] media_id={media_id} or codebook_id={codebook_id} not in project={project_id}")
                continue

            new_excerpt = Excerpt(
                project_id=project_id,
                media_id=media_id,
                codebook_id=codebook_id,
                code=code,
                subcode=subcode,
                excerpt=excerpt_text,
                explanation=explanation,
                user_id=current_user.id
            )
            db.session.add(new_excerpt)
            db.session.flush()
            excerpt_ids.append(new_excerpt.id)
            app.logger.info(f"[save_excerpt] saved excerpt id={new_excerpt.id} for media_id={media_id} code={code}")

        db.session.commit()
        app.logger.info(f"[save_excerpt] saved {len(excerpt_ids)} excerpts")
        return jsonify({
            'message': f'Successfully saved {len(excerpt_ids)} excerpts',
            'success': True,
            'excerpt_ids': excerpt_ids
        })

    except Exception as e:
        db.session.rollback()
        app.logger.exception(f"[save_excerpt] failure: {e}")
        return jsonify({'message': f'Failed to save excerpts: {str(e)}', 'success': False}), 500


# app.py (add after existing routes)
@app.route('/analyze_project', methods=['POST'])
@login_required
def analyze_project():
    try:
        data = request.get_json() or {}
        media_ids = data.get('media_ids', [])
        codebook_id = data.get('codebook_id')
        project_id = data.get('project_id')

        app.logger.info(
            f"[analyze] user={current_user.id} project={project_id} "
            f"codebook={codebook_id} media_batch={len(media_ids)} ids={media_ids}"
        )

        if not media_ids or not codebook_id or not project_id:
            app.logger.warning("[analyze] missing media_ids/codebook_id/project_id")
            return jsonify({
                'error': 'Missing required fields: media_ids, codebook_id, or project_id'
            }), 400

        project = db.session.get(Project, project_id)
        if not project:
            app.logger.warning(f"[analyze] project not found: {project_id}")
            return jsonify({'error': 'Project not found'}), 404

        if project.owner_id != current_user.id and current_user not in project.collaborators:
            app.logger.warning(f"[analyze] unauthorized user={current_user.id} project={project_id}")
            return jsonify({'error': 'Unauthorized to analyze this project'}), 403

        codebook = db.session.get(Codebook, codebook_id)
        if not codebook:
            app.logger.warning(f"[analyze] codebook not found: {codebook_id}")
            return jsonify({'error': 'Codebook not found'}), 404

        if codebook.project_id != project_id:
            app.logger.warning(f"[analyze] codebook={codebook_id} not in project={project_id}")
            return jsonify({'error': 'Codebook does not belong to this project'}), 400

        valid_media = db.session.query(Media).filter(
            Media.id.in_(media_ids),
            Media.project_id == project_id
        ).all()

        if len(valid_media) != len(media_ids):
            app.logger.warning(
                f"[analyze] some media not in project "
                f"(sent={len(media_ids)} valid={len(valid_media)})"
            )
            return jsonify({
                'error': 'One or more media files do not belong to this project'
            }), 400

        app.logger.info("[analyze] starting analysis…")
        result = analysis_service.analyze_media(
            media_ids=media_ids,
            codebook_id=codebook_id,
            project_id=project_id,
        )

        excerpts = result.get('excerpts', [])
        explanation = result.get('explanation', '')
        app.logger.info(f"[analyze] analysis done: excerpts_found={len(excerpts)}")

        for excerpt in excerpts:
            app.logger.debug(
                "[analyze] excerpt: media_id=%s code=%s subcode=%s doc=%s chunk_id=%s text=%s...",
                excerpt.get('media_id'),
                excerpt.get('code'),
                excerpt.get('subcode'),
                excerpt.get('doc'),
                excerpt.get('chunk_id'),
                (excerpt.get('text') or '')[:100]
            )

        valid_excerpts = [
            {
                'media_id': e.get('media_id'),
                'codebook_id': codebook_id,
                'code': e.get('code'),
                'subcode': e.get('subcode'),
                'excerpt': e.get('text'),        # Map 'text' to 'excerpt'
                'explanation': e.get('explanation')
            }
            for e in excerpts
            if e.get('media_id')
            and e.get('text')
            and not e.get('text').startswith("No relevant content found")
        ]

        app.logger.info(
            "[analyze] valid excerpts for saving: count=%d",
            len(valid_excerpts)
        )

        saved_excerpt_ids = []
        if valid_excerpts:
            payload = {
                'project_id': project_id,
                'excerpts': valid_excerpts,
            }
            app.logger.info(
                "[analyze] saving excerpts: count=%d",
                len(payload['excerpts'])
            )

            save_r = app.test_client().post('/save_excerpt', json=payload)

            if save_r.status_code >= 400:
                # Be defensive: response might not be JSON (e.g. CSRF 400 HTML page)
                body_json = save_r.get_json(silent=True)
                body_text = save_r.get_data(as_text=True)

                app.logger.error(
                    "[analyze] save failed status=%s body_json=%r body_text=%r",
                    save_r.status_code, body_json, body_text
                )

                message = None
                if isinstance(body_json, dict):
                    message = body_json.get('message')
                if not message:
                    message = body_text or 'Unknown error'

                return jsonify({'error': f"Save failed: {message}"}), 500

            save_json = save_r.get_json(silent=True) or {}
            saved_excerpt_ids = save_json.get('excerpt_ids', [])
            app.logger.info(
                "[analyze] save ok: %s ids=%s",
                save_json.get('message'),
                saved_excerpt_ids
            )
        else:
            app.logger.warning("[analyze] no valid excerpts to save")
            explanation += (
                "\nNo relevant excerpts found in the document for the selected codebook. "
                "Consider revising the codebook or selecting a different document."
            )

        return jsonify({
            'excerpts': excerpts,
            'explanation': explanation,
            'saved_excerpt_ids': saved_excerpt_ids,
        })

    except Exception as e:
        app.logger.exception(f"[analyze] failure: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500



# Add this after existing routes (e.g., after @app.route('/analyze_project', methods=['POST']) block)
@app.route('/analyze_policy', methods=['POST'])
@login_required  # Optional: require login if needed; remove if public
def analyze_policy():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        if not media_allowed_file(file.filename):  # Reuse existing helper
            return jsonify({'error': 'Invalid file type. Only PDF/DOCX allowed'}), 400

        # Save temp file (reuse upload logic)
        filename = secure_filename(file.filename)
        
        # First check if file already exists by filename (quick check)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(filepath):
            flash(f'File already exists: {filename}. Using existing file for analysis.', 'warning')
            return redirect(request.referrer or url_for('dashboard'))

        file.save(filepath)

        # Extract text (reuse DocumentProcessor without modifying it)
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            policy_text = doc_processor.extract_text_from_pdf(filepath)
        elif ext == '.docx':
            policy_text = doc_processor.extract_text_from_docx(filepath)
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        if not policy_text.strip():
            return jsonify({'error': 'No text extracted from document'}), 400

        # Analyze (new service)
        analyzer = PolicyAnalyzer()
        keywords = analyzer.derive_keywords(policy_text)
        evidence = analyzer.search_evidence(keywords)
        suggestions = analyzer.generate_suggestions(policy_text, evidence)

        # Cleanup temp file
        os.remove(filepath)

        app.logger.info(f"[analyze_policy] success: keywords={keywords}, evidence={len(evidence)}, suggestions generated")

        return jsonify({
            'success': True,
            'suggestions': suggestions,
            'evidence': evidence,  # For displaying links
            'keywords': keywords  # Optional: show derived keywords
        })

    except Exception as e:
        app.logger.exception(f"[analyze_policy] failure: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

# Add this route after existing ones
@app.route('/live_suggest', methods=['POST'])
@login_required  # Optional: for authenticated users
def live_suggest():
    try:
        data = request.get_json() or {}
        context = data.get('context', '')  # From JS: Now section + prior
        product_type = data.get('product_type', '')

        if not context.strip():
            return jsonify({'error': 'No context provided'}), 400

        # Reuse existing analyzer for evidence
        analyzer = PolicyAnalyzer()

        prompt_keywords = f"{product_type} policy {context}" if product_type else context
        # In live_suggest route, replace the evidence line:
        try:
            keywords = analyzer.derive_keywords(prompt_keywords)
            evidence = analyzer.search_evidence(keywords)
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 429:
                app.logger.warning("[live_suggest] Google API rate limit; using no evidence")
                evidence = []
                keywords = []  # Or fallback keywords
            else:
                raise
        except Exception as e:
            app.logger.error(f"[live_suggest] Search error: {e}")
            evidence = []
            keywords = []

        # ... (prompt uses evidence as-is; LLM still runs)


        # LLM for suggestions (reuse policy_builder.llm)
        prompt = f"""
        You are a real-time policy writing assistant. Current Section Context: "{context}"
        Product Type: {product_type or 'General Policy Document'}

        RELEVANT EVIDENCE (prioritized academics first):
        {chr(10).join([f"- {r['title']} ({r['source']}): {r['snippet'][:150]}..." for r in evidence])}

        TASK: Provide 2-3 concise suggestions to continue in the CURRENT SECTION ONLY (e.g., autocomplete sentence or paragraph). 
        Use prior context for flow, but focus on this section's content. Evidence-based, professional for {product_type.lower()} drafting. 
        Format as JSON array: ["Suggestion 1", "Suggestion 2", "Suggestion 3"].
        """
        try:
            response = analyzer.policy_builder.llm.invoke(prompt)
            # Simple parse (assume LLM outputs clean JSON; enhance with json.loads if needed)
            suggestions = [s.strip() for s in response.split('\n') if s.strip() and not s.startswith('TASK:')]
        except:
            suggestions = ["Based on evidence, consider adding: 'Implementation should include monitoring via WHO guidelines.'"]

        app.logger.info(f"[live_suggest] generated {len(suggestions)} suggestions for context len={len(context)}")

        return jsonify({
            'success': True,
            'suggestions': suggestions[:3],  # Limit to 3
            'evidence': evidence[:5],  # Top 5 for display
            'keywords': keywords
        })

    except Exception as e:
        app.logger.exception(f"[live_suggest] failure: {e}")
        return jsonify({'error': f'Suggestion failed: {str(e)}'}), 500


# Add this route after @app.route('/live_suggest', methods=['POST'])
@app.route('/get_structure', methods=['POST'])
@login_required  # Consistent with live_writing
def get_structure():
    try:
        data = request.get_json() or {}
        product_type = data.get('product_type', '').strip()
        if not product_type:
            return jsonify({'error': 'No product type provided'}), 400

        # Validate against allowed types
        allowed_types = [
            'Policy Briefs', 'Policy Frameworks', 'Legislative Instruments / Policy Framework Laws',
            'Strategic Plans / Implementation Strategies', 'White Papers / Green Papers'
        ]
        if product_type not in allowed_types:
            return jsonify({'error': 'Invalid product type'}), 400

        # Get general description for prompting (hardcoded for now; could DB later)
        descriptions = {
            'Policy Briefs': 'Concise, targeted documents, usually 2–6 pages, that summarize evidence, outline policy options, and make actionable recommendations for decision-makers.',
            'Policy Frameworks': 'High-level guiding documents that set the vision, principles, and strategic direction for addressing an issue, e.g., a National Health Policy Framework.',
            'Legislative Instruments / Policy Framework Laws': 'Formal legal instruments that establish binding obligations and provide the legal foundation for policy implementation, e.g., Climate Change Act, Data Protection Act.',
            'Strategic Plans / Implementation Strategies': 'Detailed roadmaps translating policy frameworks into time-bound actions, with defined priorities, objectives, and measurable indicators, e.g., National Strategic Plan for HIV/AIDS.',
            'White Papers / Green Papers': 'Government discussion or consultative documents that frame an issue, outline possible approaches, and invite stakeholder input before final policy adoption.'
        }
        description = descriptions.get(product_type, '')

        # Reuse analyzer for evidence (use product_type as initial keyword for relevance)
        analyzer = PolicyAnalyzer()
        keywords = analyzer.derive_keywords(product_type + ' ' + description)  # Derive from type + desc
        evidence = analyzer.search_evidence(keywords)

        # LLM prompt for structure (evidence-based outline)
        evidence_summary = "\n".join([f"- {r['title']} ({r['source']}): {r['snippet'][:150]}..." for r in evidence])
        prompt = f"""
        You are a policy drafting assistant. For a {product_type}: {description}

        RELEVANT EVIDENCE (prioritized academics first):
        {evidence_summary or "No evidence found; use general best practices."}

        TASK:
        - Generate a pre-drafted structure/outline as a starting template.
        - Use Markdown-like formatting: # Title, ## Section, ### Subsection, with placeholder text (e.g., "Insert evidence here: [cite source]").
        - Include 4-8 main sections based on standard structures, infused with evidence where relevant.
        - Keep professional, concise; end with placeholders for user input.
        - Output ONLY the formatted outline text.
        """

        try:
            outline = analyzer.policy_builder.llm.invoke(prompt)
            app.logger.info(f"[get_structure] generated outline for {product_type}: {len(evidence)} evidence items")
            return jsonify({
                'success': True,
                'outline': outline,
                'evidence': evidence[:5],  # Initial evidence for display
                'keywords': keywords
            })
        except Exception as e:
            app.logger.exception(f"[get_structure] LLM error: {e}")
            return jsonify({'error': 'Failed to generate structure'}), 500

    except Exception as e:
        app.logger.exception(f"[get_structure] failure: {e}")
        return jsonify({'error': f'Structure generation failed: {str(e)}'}), 500

# -----------------------------------------------------------------------------
# User Management (Admin Only)
# -----------------------------------------------------------------------------
@app.route('/admin/users')
@login_required
def list_users():
    if not can_edit_all():
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    users = User.query.all()
    return render_template('admin/users.html', users=users)

@app.route('/admin/users/create', methods=['POST'])
@login_required
def create_user():
    if not can_edit_all():
        abort(403)
    
    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')
    is_admin = 'is_admin' in request.form
    is_view_only_admin = 'is_view_only_admin' in request.form
    
    if not username or not email or not password:
        flash('Username, email, and password are required.', 'danger')
        return redirect(url_for('list_users'))
        
    if User.query.filter((User.username == username) | (User.email == email)).first():
        flash('Username or email already exists.', 'danger')
        return redirect(url_for('list_users'))
        
    new_user = User(username=username, email=email, is_admin=is_admin, is_view_only_admin=is_view_only_admin)
    new_user.set_password(password)
    
    try:
        db.session.add(new_user)
        db.session.commit()
        flash(f'User {username} created successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error creating user: {str(e)}', 'danger')
        
    return redirect(url_for('list_users'))

@app.route('/admin/users/edit/<int:user_id>', methods=['POST'])
@login_required
def edit_user(user_id):
    app.logger.info(f"Attempting to edit user {user_id}")
    app.logger.info(f"Form data: {request.form}")
    if not can_edit_all():
        app.logger.warning("Access denied for edit_user")
        abort(403)
    user = User.query.get_or_404(user_id)
    
    # Update roles
    user.is_admin = 'is_admin' in request.form
    user.is_view_only_admin = 'is_view_only_admin' in request.form
    
    # Update password if provided
    password = request.form.get('password')
    if password and password.strip():
        user.set_password(password.strip())
        
    try:
        db.session.commit()
        flash(f'User {user.username} updated successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error updating user: {str(e)}', 'danger')
        
    return redirect(url_for('list_users'))

@app.route('/admin/users/delete/<int:user_id>', methods=['POST'])
@login_required
def delete_user(user_id):
    if not can_edit_all():
        abort(403)
    if user_id == current_user.id:
        flash('You cannot delete yourself.', 'danger')
        return redirect(url_for('list_users'))
        
    user = User.query.get_or_404(user_id)
    try:
        db.session.delete(user)
        db.session.commit()
        flash(f'User {user.username} deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting user: {str(e)}', 'danger')
        
    return redirect(url_for('list_users'))
@app.route('/live_writing')
@login_required  # Optional: for authenticated users; remove if public
def live_writing():
    """
    Renders the Live Writing Assisted by AI page.
    """
    try:
        return render_template('research/live_writing.html')
    except Exception as e:
        app.logger.exception(f"[live_writing] template render failure: {e}")
        abort(500)


# -----------------------------------------------------------------------------
# Home
# -----------------------------------------------------------------------------
@app.route('/')
def index():
    return render_template('index.html')

# -----------------------------------------------------------------------------
# Entrypoint
# -----------------------------------------------------------------------------
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
